import gradio as gr
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import subprocess
import threading
import time
import os
import psutil
import GPUtil
import re
from datetime import datetime

# ============================================================================
# CONFIGURATION & DEFAULTS
# ============================================================================

DEFAULT_PROMPT = (
    "You are a professional English-Arabic translator. "
    "Translate the following text from English to Arabic in a refined, sophisticated, and professional book-author style. "
    "Ensure that your translation is clear and does not include any notes, disclaimers, or commentary about the original text's quality. "
    "Replace numerical figures with their corresponding words in Arabic, including large numbers, decimals, and scientific values. "
    "Remove any citation numbers that appear at the end of sentences. "
    "Translate all words, including scientific terms. "
    "Only provide the translated text.\n\n"
)

DEBUG_LOG = os.path.join(os.getcwd(), 'translator_debug.log')

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def write_log(msg: str):
    """Write timestamped log message to debug file"""
    ts = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    try:
        with open(DEBUG_LOG, 'a', encoding='utf-8') as f:
            f.write(f"[{ts}] {msg}\n")
    except Exception:
        pass

def sanitize_output(text: str) -> str:
    """Remove internal reasoning phrases from model output"""
    try:
        if not text:
            return text
        
        lines = text.splitlines()
        filtered = []
        thought_pattern = re.compile(
            r"\b(thought|thinking|reasoning|analysis|chain[- ]of[- ]thought)\b", 
            re.IGNORECASE
        )
        
        for ln in lines:
            if thought_pattern.search(ln):
                continue
            if ln.strip().lower() in ("thinking...", "thought for a moment", "thinking"):
                continue
            filtered.append(ln)
        
        cleaned = "\n".join(filtered).strip()
        
        # Extract Arabic text if present
        arabic_runs = re.findall(
            r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]+(?:[^\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]+[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]+)*", 
            cleaned
        )
        if arabic_runs:
            return "\n".join(r.strip() for r in arabic_runs).strip()
        
        return cleaned
    except Exception:
        return text

def get_system_stats():
    """Get current system resource usage"""
    cpu = psutil.cpu_percent(interval=0.1)
    mem = psutil.virtual_memory().percent
    
    try:
        gpus = GPUtil.getGPUs()
        if gpus:
            gpu_stats = f"🎮 GPU: {gpus[0].load*100:.1f}% | VRAM: {gpus[0].memoryUtil*100:.1f}%"
        else:
            gpu_stats = "❌ No GPU detected"
    except Exception as e:
        gpu_stats = f"⚠️ GPU error: {str(e)[:30]}"
    
    return f"🖥️ CPU: {cpu:.1f}% | 🧠 RAM: {mem:.1f}% | {gpu_stats}"

def list_llms():
    """List available Ollama models"""
    try:
        result = subprocess.run(["ollama", "list"], capture_output=True, text=True)
        models = [
            line.split()[0] for line in result.stdout.splitlines() 
            if line.strip() and not line.startswith("NAME")
        ]
        return models if models else ["command-r7b-arabiclatest"]
    except Exception:
        return ["command-r7b-arabiclatest"]

def calculate_optimal_batch_size(file_path):
    """Calculate optimal batch size based on document size"""
    try:
        text = read_text_from_file(file_path)
        total_chars = len(text)
    except Exception:
        # Fallback: assume medium size
        total_chars = 20000
    
    # Smart batch sizing: 1000 chars for small docs, scales up for larger docs
    if total_chars < 5000:
        batch_size = 1000
    elif total_chars < 50000:
        batch_size = 2000
    elif total_chars < 200000:
        batch_size = 3000
    else:
        batch_size = 4000
    
    write_log(f"Document size: {total_chars} chars | Optimal batch size: {batch_size}")
    return batch_size

# ============================================================================
# DOCUMENT HANDLING
# ============================================================================

def read_docx(file_path):
    """Read text from DOCX file"""
    doc = Document(file_path)
    return "\n".join(para.text for para in doc.paragraphs)


def read_text_from_file(file_path):
    """Read text from various file types: .docx, .pdf, .epub, .txt, .html

    Uses optional dependencies when available and falls back gracefully.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.docx':
        return read_docx(file_path)

    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    if ext == '.html' or ext == '.htm':
        try:
            import importlib
            bs4 = importlib.import_module('bs4')
            BeautifulSoup = bs4.BeautifulSoup
        except Exception:
            # simple strip tags fallback
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return re.sub(r'<[^>]+>', '', f.read())
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            soup = BeautifulSoup(f, 'html.parser')
            return soup.get_text(separator='\n')

    if ext == '.pdf':
        # Try PyMuPDF (fitz) first, then pdfminer as fallback
        try:
            import importlib
            fitz = importlib.import_module('fitz')
            doc = fitz.open(file_path)
            txt = []
            for page in doc:
                txt.append(page.get_text())
            return '\n'.join(txt)
        except Exception:
            try:
                # pdfminer.six
                import importlib
                pdfminer = importlib.import_module('pdfminer.high_level')
                extract_text = getattr(pdfminer, 'extract_text')
                return extract_text(file_path)
            except Exception:
                # last resort: empty
                return ''

    if ext == '.epub':
        try:
            import importlib
            epub = importlib.import_module('ebooklib.epub')
            bs4 = importlib.import_module('bs4')
            BeautifulSoup = bs4.BeautifulSoup
            book = epub.read_epub(file_path)
            texts = []
            for item in book.get_items():
                if item.get_type() == epub.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_body_content(), 'html.parser')
                    texts.append(soup.get_text(separator='\n'))
            return '\n'.join(texts)
        except Exception:
            return ''

    # Unknown extension: try reading as text
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception:
        return ''

def write_docx(file_path, text, direction="rtl", alignment="right"):
    """Write text to DOCX file with formatting options"""
    doc = Document()
    
    # Alignment mapping
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    
    alignment_enum = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    
    for para_text in text.split('\n'):
        para = doc.add_paragraph(para_text)
        para.alignment = alignment_enum
        
        # Set text direction (RTL for Arabic)
        if direction == "rtl":
            para.paragraph_format.right_to_left = True
        
        # Set font properties
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
    
    doc.save(file_path)

# ============================================================================
# TRANSLATION STATE MANAGEMENT
# ============================================================================

class TranslationState:
    """Thread-safe translation state manager"""
    
    def __init__(self):
        self.lock = threading.Lock()
        self.reset()
    
    def reset(self):
        """Reset all state variables"""
        with self.lock:
            self.paused = False
            self.stopped = False
            self.running = False
            self.current_chunk = 0
            self.total_chunks = 0
            self.translated_chunks = []
            self.live_feed = ""
            self.status_message = ""
            self.start_time = None
            self.prompt_template = None
            self.prompt_mode = 'append'
            self.think_enabled = True
            self.concise_mode = False
            self.filter_thoughts = False
            self.doc_direction = "rtl"
            self.doc_alignment = "right"
    
    def get_progress(self):
        """Get current progress percentage"""
        with self.lock:
            if self.total_chunks == 0:
                return 0
            return int(100 * self.current_chunk / self.total_chunks)
    
    def get_status(self):
        """Get formatted status message with stats"""
        with self.lock:
            if not self.running:
                return self.status_message
            
            elapsed = time.time() - self.start_time if self.start_time else 0
            chunks_left = self.total_chunks - self.current_chunk
            
            if self.current_chunk > 0:
                avg_time = elapsed / self.current_chunk
                eta = avg_time * chunks_left
                eta_str = f"⏳ ETA: {int(eta)}s"
            else:
                eta_str = "⏳ Calculating..."
            
            sys_stats = get_system_stats()
            
            if self.paused:
                return f"⏸️ PAUSED - Chunk {self.current_chunk}/{self.total_chunks}\n{sys_stats}"
            
            return (
                f"🔄 Translating chunk {self.current_chunk}/{self.total_chunks} | "
                f"{eta_str}\n{sys_stats}"
            )
    
    def get_live_feed(self):
        """Get current live feed"""
        with self.lock:
            return self.live_feed

translation_state = TranslationState()

# ============================================================================
# OLLAMA TRANSLATION
# ============================================================================

def translate_with_ollama(text, model="command-r7b-arabiclatest", temperature=0.0):
    """Call Ollama to translate text chunk"""
    
    # Build prompt
    prompt_template = translation_state.prompt_template
    prompt_mode = translation_state.prompt_mode
    
    if prompt_template:
        if "{text}" in prompt_template:
            prompt = prompt_template.format(text=text)
        else:
            if prompt_mode == "replace":
                prompt = prompt_template + "\n\n" + text
            else:
                prompt = DEFAULT_PROMPT + "\n" + prompt_template + "\n\n" + text + "\n\nArabic Translation:"
    else:
        prompt = DEFAULT_PROMPT + text + "\n\nArabic Translation:"
    
    # Add concise instruction if enabled
    if translation_state.concise_mode:
        prompt = "Respond concisely; minimal reasoning.\n\n" + prompt
    
    # Build command
    cmd = ["ollama", "run", model]
    if not translation_state.think_enabled:
        cmd.append("--think=false")
    cmd.append(prompt)
    
    write_log(f"Ollama call: model={model}, think={translation_state.think_enabled}, prompt_len={len(prompt)}")
    
    # Execute
    try:
        result = subprocess.run(
            cmd, 
            capture_output=True, 
            text=True, 
            encoding='utf-8', 
            errors='replace',
            timeout=300
        )
    except subprocess.TimeoutExpired:
        return "⚠️ Translation timeout (5 minutes exceeded)"
    except Exception as e:
        return f"❌ Exception: {e}"
    
    if result.returncode != 0:
        out = result.stderr or result.stdout
        return "❌ Error: " + (out.strip() if out else "unknown error")
    
    raw = (result.stdout or "").strip()
    
    # Apply thought filtering if enabled
    if translation_state.filter_thoughts and raw:
        raw = sanitize_output(raw)
    
    return raw if raw else "⚠️ No output received from Ollama."

# ============================================================================
# TRANSLATION WORKER
# ============================================================================

def translation_worker(input_file, model, batch_size):
    """Background worker for translation process"""
    try:
        write_log(f"Worker started: file={input_file.name}, model={model}, batch={batch_size}")

        # Read document (supports multiple formats)
        text = read_text_from_file(input_file.name)
        total_chars = len(text)

        # Calculate chunks
        total_chunks = (total_chars + batch_size - 1) // batch_size

        with translation_state.lock:
            translation_state.running = True
            translation_state.total_chunks = total_chunks
            translation_state.current_chunk = 0
            translation_state.translated_chunks = []
            translation_state.live_feed = ""
            translation_state.start_time = time.time()

        # Process chunks
        for i in range(0, total_chars, batch_size):
            # Check for stop/pause
            with translation_state.lock:
                if translation_state.stopped:
                    translation_state.status_message = "🛑 Translation stopped by user"
                    break

                while translation_state.paused:
                    time.sleep(0.1)

            # Extract chunk
            chunk = text[i:i+batch_size]
            chunk_num = i // batch_size + 1

            write_log(f"Processing chunk {chunk_num}/{total_chunks}, size={len(chunk)}")

            # Translate
            translated = translate_with_ollama(chunk, model=model)

            # Store results - UPDATE IN REAL TIME
            with translation_state.lock:
                translation_state.translated_chunks.append(translated)
                translation_state.current_chunk = chunk_num
                translation_state.live_feed += translated + "\n"

        # Finalize
        with translation_state.lock:
            translation_state.running = False

            if not translation_state.stopped:
                # Save output
                output_dir = os.path.dirname(input_file.name)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = os.path.join(output_dir, f"translated_{timestamp}.docx")

                full_translation = "\n".join(translation_state.translated_chunks)
                write_docx(output_path, full_translation, direction=translation_state.doc_direction,
                          alignment=translation_state.doc_alignment)

                elapsed = time.time() - translation_state.start_time
                translation_state.status_message = (
                    f"✅ Translation complete!\n"
                    f"⏱️ Time: {int(elapsed)}s | 📄 Chunks: {total_chunks} | "
                    f"📊 {get_system_stats()}\n"
                    f"💾 Saved to: {os.path.basename(output_path)}"
                )

                write_log(f"Translation completed: {output_path}")

    except Exception as e:
        write_log(f"Worker error: {e}")
        with translation_state.lock:
            translation_state.running = False
            translation_state.status_message = f"❌ Error: {e}"

# ============================================================================
# GRADIO INTERFACE FUNCTIONS
# ============================================================================

def start_translation(input_file, model, prompt_template, prompt_mode, 
                     disable_think, enable_concise, enable_filter):
    """Initialize and start translation process"""
    
    if not input_file:
        return "❌ Please upload a DOCX file first", gr.update(value=0), "", gr.update(interactive=True)
    
    # Calculate optimal batch size
    batch_size = calculate_optimal_batch_size(input_file.name)
    
    # Reset state
    translation_state.reset()
    
    # Configure options
    with translation_state.lock:
        translation_state.prompt_template = prompt_template if prompt_template.strip() else None
        translation_state.prompt_mode = prompt_mode
        translation_state.think_enabled = not disable_think
        translation_state.concise_mode = enable_concise
        translation_state.filter_thoughts = enable_filter
    
    # Start worker thread
    thread = threading.Thread(
        target=translation_worker,
        args=(input_file, model, batch_size),
        daemon=True
    )
    thread.start()
    
    write_log(f"Translation started: think={not disable_think}, concise={enable_concise}, filter={enable_filter}, batch_size={batch_size}")
    
    return "🚀 Translation started...", gr.update(value=0), "⏳ Processing...", gr.update(interactive=False)

def poll_progress():
    """Generator that continuously polls translation progress"""
    last_feed = ""
    while True:
        progress = translation_state.get_progress()
        status = translation_state.get_status()
        live_feed = translation_state.get_live_feed()
        is_running = translation_state.running
        
        # Only yield if something changed or still running
        if live_feed != last_feed or is_running:
            last_feed = live_feed
            yield gr.update(value=progress), gr.update(value=status), gr.update(value=live_feed), gr.update(interactive=not is_running)
        
        # Stop polling when translation is done
        if not is_running:
            break
        
        time.sleep(0.3)

def pause_toggle():
    """Toggle pause state"""
    with translation_state.lock:
        translation_state.paused = not translation_state.paused
        state = "⏸️ Paused" if translation_state.paused else "▶️ Resumed"
    
    write_log(f"Pause toggled: {state}")
    
    # Return current state
    progress = translation_state.get_progress()
    status = translation_state.get_status()
    live_feed = translation_state.get_live_feed()
    is_running = translation_state.running
    
    return gr.update(value=progress), gr.update(value=status), gr.update(value=live_feed), gr.update(interactive=not is_running)

def stop_translation():
    """Stop translation process"""
    with translation_state.lock:
        translation_state.stopped = True
    
    write_log("Translation stopped by user")
    time.sleep(0.5)
    
    progress = translation_state.get_progress()
    status = translation_state.get_status()
    live_feed = translation_state.get_live_feed()
    
    return gr.update(value=progress), gr.update(value=status), gr.update(value=live_feed), gr.update(interactive=True)

def export_and_download():
    """Export current translation to DOCX"""
    with translation_state.lock:
        if not translation_state.translated_chunks:
            return None, "❌ No translation to export"
        
        output_path = os.path.join(os.getcwd(), f"translation_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        full_text = "\n".join(translation_state.translated_chunks)
        
        write_docx(output_path, full_text, direction=translation_state.doc_direction, alignment=translation_state.doc_alignment)
    
    write_log(f"Export: {output_path}")
    return output_path, f"✅ Ready to download!"

# ============================================================================
# GRADIO UI LAYOUT
# ============================================================================

with gr.Blocks(title="📚 Tajruba PBT ", theme=gr.themes.Soft()) as demo:

    gr.Markdown("# 📚 Tajreba ")
    gr.Markdown("Professional Book Translator , Translate DOCX documents from English to Arabic using local LLM models via Ollama")

    # Inject custom CSS to better match the reference UI and stable element IDs
    gr.HTML("""
    <style>
      body { background: #091021; color: #e6eef6; font-family: Inter, ui-sans-serif, system-ui, -apple-system, 'Segoe UI', Roboto, 'Helvetica Neue', Arial; }
      .gradio-container { background: transparent; }
      /* Status and live output boxes */
      #status_box .input, #live_output .input { background: #071026 !important; color: #d9eefb !important; border: 1px solid #123243 !important; }
      /* Buttons */
      #translate_btn button, #export_btn button { background: linear-gradient(90deg,#4f46e5,#06b6d4) !important; color: #fff !important; border-radius: 8px; }
      #pause_btn button, #stop_btn button { background: #0b7285 !important; color: #fff !important; border-radius: 8px; }
      /* Progress slider */
      #progress_bar .gr-slider { background: linear-gradient(90deg,#0ea5a4,#6366f1) !important; }
      /* Live output look */
      #live_output textarea { background: #041427 !important; color: #e6f3ff !important; border-radius: 6px; }
      /* Smaller accents */
      .gr-block-title { color: #cfeefe !important; }
    </style>
    """)

    with gr.Row():
        # LEFT PANEL
        with gr.Column(scale=1):
            # File input (accept multiple document formats)
            input_file = gr.File(
                label="📂 Upload Document (.docx, .pdf, .epub, .txt, .html)",
                file_types=[".docx", ".pdf", ".epub", ".txt", ".html"],
                type="filepath"
            )

            # Model selection
            llm_choices = list_llms()
            default_llm = "command-r7b-arabiclatest" if "command-r7b-arabiclatest" in llm_choices else (llm_choices[0] if llm_choices else "command-r7b-arabiclatest")
            model_dropdown = gr.Dropdown(
                choices=llm_choices,
                value=default_llm,
                label="🤖 Select LLM Model"
            )

            # Control buttons
            with gr.Row():
                translate_btn = gr.Button("🚀 Start Translation", variant="primary", size="lg", scale=1, elem_id="translate_btn")

            with gr.Row():
                pause_btn = gr.Button("⏸️ Pause", size="lg", scale=1, elem_id="pause_btn")
                stop_btn = gr.Button("🛑 Stop", size="lg", scale=1, elem_id="stop_btn")

            export_download_btn = gr.Button("💾 Export & Download", variant="primary", size="lg", elem_id="export_btn")

            # Advanced Options
            with gr.Accordion("⚙️ Advanced Options", open=False):
                prompt_template = gr.Textbox(
                    label="✍️ Custom Prompt",
                    value=DEFAULT_PROMPT,
                    lines=4,
                    placeholder="Leave empty to use default prompt"
                )

                prompt_mode = gr.Radio(
                    choices=["append", "replace"],
                    value="append",
                    label="Prompt Mode"
                )

                disable_think = gr.Checkbox(
                    label="🧠 Disable reasoning",
                    value=False
                )

                enable_concise = gr.Checkbox(
                    label="✂️ Concise mode",
                    value=False
                )

                enable_filter = gr.Checkbox(
                    label="🔍 Filter thoughts",
                    value=True
                )

            # removed Output Formatting section (language/direction/alignment)
            doc_direction = gr.State("rtl")
            doc_alignment = gr.State("right")
            output_file = gr.File(label="📥 Download", interactive=False)

        # RIGHT PANEL
        with gr.Column(scale=2):
            # Progress section
            gr.Markdown("### 📊 Translation Progress")

            progress_bar = gr.Slider(
                minimum=0,
                maximum=100,
                value=0,
                label="Progress",
                interactive=False,
                elem_id="progress_bar",
                show_label=False
            )

            status_box = gr.Textbox(
                label="Status & System Info",
                value="Ready to translate",
                lines=3,
                interactive=False,
                elem_id="status_box"
            )

            # Live output section
            gr.Markdown("### 📝 Live Translation Output")
            live_output = gr.Textbox(
                label="",
                lines=20,
                max_lines=30,
                interactive=False,
                show_label=False,
                placeholder="Translation output will appear here in real-time...",
                elem_id="live_output"
            )

    # ========================================================================
    # EVENT BINDINGS
    # ========================================================================

    # (Output formatting controls removed)

    # Start translation with streaming updates
    translate_btn.click(
        fn=start_translation,
        inputs=[
            input_file,
            model_dropdown,
            prompt_template,
            prompt_mode,
            disable_think,
            enable_concise,
            enable_filter
        ],
        outputs=[status_box, progress_bar, live_output, translate_btn]
    ).then(
        fn=poll_progress,
        outputs=[progress_bar, status_box, live_output, translate_btn]
    )

    # Control buttons
    pause_btn.click(
        fn=pause_toggle,
        outputs=[progress_bar, status_box, live_output, translate_btn]
    )

    stop_btn.click(
        fn=stop_translation,
        outputs=[progress_bar, status_box, live_output, translate_btn]
    )

    # Export & Download
    export_download_btn.click(
        fn=export_and_download,
        outputs=[output_file, status_box]
    )

# ============================================================================
# LAUNCH
# ============================================================================

if __name__ == "__main__":
    write_log("=" * 60)
    write_log("Translator UI starting...")
    demo.launch(server_name="127.0.0.1", server_port=7860)

# Ensure required packages are installed
try:
    import importlib
    import sys
    import fitz
    import pdfminer.high_level
    import bs4
    import ebooklib.epub
except ImportError:
    print("Installing required packages: pymupdf, pdfminer.six, beautifulsoup4, ebooklib")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "pymupdf", "pdfminer.six", "beautifulsoup4", "ebooklib"])