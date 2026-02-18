'''
Author:     Ayoub Baydy
LinkedIn:   https://www.linkedin.com/in/ayoubaydy/

Copyright (C) 2024 Ayoub Baydy

GitHub:     https://github.com/ayoubaydy/Tajreba

version:    18.02.26.1.00
'''

import os
import time
import threading
import subprocess
import re
from datetime import datetime
from flask import Flask, request, jsonify, send_file, render_template
from werkzeug.utils import secure_filename
import psutil
try:
    import GPUtil
except ImportError:
    GPUtil = None

# Import document handling libraries
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Initialize Flask App
# We set static_folder and template_folder relative to this file
app = Flask(__name__, static_folder='static', template_folder='templates')

# Configuration
# Uploads go to project_root/uploads to keep src clean
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

# ============================================================================
# CONFIGURATION & DEFAULTS
# ============================================================================

DEFAULT_PROMPT = (
    "You are a professional English-Arabic translator. "
    "Translate the following text from English to Arabic in a refined, sophisticated, and professional book-author style. "
    "Ensure that your translation is clear and does not include any notes, disclaimers, or commentary about the original text's quality. "
    "Replace numerical figures with their corresponding words in Arabic, including large numbers, decimals, and scientific values. "
    "Remove any citation numbers that appear at the end of sentences. "
    "Translate all words, including scientific terms. "
    "Only provide the translated text.\n\n"
)

# Debug log in project root
DEBUG_LOG = os.path.join(BASE_DIR, 'translator_debug.log')

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def write_log(msg: str):
    """Write timestamped log message to debug file"""
    ts = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    try:
        with open(DEBUG_LOG, 'a', encoding='utf-8') as f:
            f.write(f"[{ts}] {msg}\n")
    except Exception:
        pass

def sanitize_output(text: str) -> str:
    """Remove internal reasoning phrases from model output"""
    try:
        if not text:
            return text
        
        lines = text.splitlines()
        filtered = []
        thought_pattern = re.compile(
            r"\b(thought|thinking|reasoning|analysis|chain[- ]of[- ]thought)\b", 
            re.IGNORECASE
        )
        
        for ln in lines:
            if thought_pattern.search(ln):
                continue
            if ln.strip().lower() in ("thinking...", "thought for a moment", "thinking"):
                continue
            filtered.append(ln)
        
        cleaned = "\n".join(filtered).strip()
        
        # Extract Arabic text if present. Simple heuristic.
        arabic_runs = re.findall(
            r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]+(?:[^\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]+[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]+)*", 
            cleaned
        )
        if arabic_runs:
            return "\n".join(r.strip() for r in arabic_runs).strip()
        
        return cleaned
    except Exception:
        return text

def get_system_stats():
    """Get current system resource usage"""
    cpu = psutil.cpu_percent(interval=None)
    mem = psutil.virtual_memory().percent
    
    try:
        if GPUtil:
            gpus = GPUtil.getGPUs()
            if gpus:
                gpu_stats = f"🎮 GPU: {gpus[0].load*100:.1f}% | VRAM: {gpus[0].memoryUtil*100:.1f}%"
            else:
                gpu_stats = "❌ No GPU"
        else:
            gpu_stats = ""
    except Exception:
        gpu_stats = ""
    
    return f"CPU: {cpu:.1f}% | RAM: {mem:.1f}% {gpu_stats}"

def calculate_optimal_batch_size(file_path):
    """Calculate optimal batch size based on document size"""
    try:
        text = read_text_from_file(file_path)
        total_chars = len(text)
    except Exception:
        total_chars = 20000
    
    # Smart batch sizing
    if total_chars < 5000:
        batch_size = 1000
    elif total_chars < 50000:
        batch_size = 2000
    elif total_chars < 200000:
        batch_size = 3000
    else:
        batch_size = 4000
    
    return batch_size

# ============================================================================
# DOCUMENT HANDLING
# ============================================================================

def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join(para.text for para in doc.paragraphs)

def read_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.docx':
        return read_docx(file_path)

    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    if ext == '.html' or ext == '.htm':
        try:
            import bs4
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = bs4.BeautifulSoup(f, 'html.parser')
                return soup.get_text(separator='\n')
        except Exception:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return re.sub(r'<[^>]+>', '', f.read())

    if ext == '.pdf':
        try:
            import fitz
            doc = fitz.open(file_path)
            txt = []
            for page in doc:
                txt.append(page.get_text())
            return '\n'.join(txt)
        except Exception:
             return '' 

    if ext == '.epub':
        try:
            import ebooklib.epub as epub
            import bs4
            book = epub.read_epub(file_path)
            texts = []
            for item in book.get_items():
                if item.get_type() == epub.ITEM_DOCUMENT:
                    soup = bs4.BeautifulSoup(item.get_body_content(), 'html.parser')
                    texts.append(soup.get_text(separator='\n'))
            return '\n'.join(texts)
        except Exception:
            return ''

    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception:
        return ''

def write_docx(file_path, text, direction="rtl", alignment="right"):
    doc = Document()
    
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    
    alignment_enum = align_map.get(alignment, WD_ALIGN_PARAGRAPH.RIGHT)
    
    # Split by double newline to preserve paragraph structure somewhat
    paragraphs = text.split('\n')
    
    for para_text in paragraphs:
        if not para_text.strip():
            continue
        para = doc.add_paragraph(para_text)
        para.alignment = alignment_enum
        
        if direction == "rtl":
            para.paragraph_format.right_to_left = True
        
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            # Ensure RTL context for runs if needed
            if direction == "rtl":
                run.font.rtl = True
    
    doc.save(file_path)

# ============================================================================
# TRANSLATION STATE
# ============================================================================

class TranslationState:
    def __init__(self):
        self.lock = threading.Lock()
        self.reset()
    
    def reset(self):
        with self.lock:
            self.paused = False
            self.stopped = False
            self.running = False
            self.current_chunk = 0
            self.total_chunks = 0
            self.translated_chunks = []
            self.live_feed = ""
            self.status_message = ""
            self.start_time = None
            self.error = None
            self.output_file = None
            self.current_config = {}

state = TranslationState()

# ============================================================================
# LOGIC
# ============================================================================

def translate_with_ollama(text, model="command-r7b-arabiclatest", temperature=0.3):
    """Call Ollama to translate text chunk"""
    prompt = DEFAULT_PROMPT + text + "\n\nArabic Translation:"
    
    cmd = ["ollama", "run", model]
    
    try:
        result = subprocess.run(
            cmd, 
            input=prompt,
            capture_output=True, 
            text=True, 
            encoding='utf-8', 
            errors='replace',
            timeout=300
        )
    except Exception as e:
        return f"[Error: {e}]"
    
    if result.returncode != 0:
        return f"[Error: {result.stderr or 'Unknown error'}]"
    
    return sanitize_output(result.stdout.strip())

def worker_thread(filepath, config):
    write_log(f"Worker started for {filepath}")
    
    try:
        # 1. Read File
        text = read_text_from_file(filepath)
        total_chars = len(text)
        
        # 2. Chunking
        batch_size = calculate_optimal_batch_size(filepath)
        total_chunks = (total_chars + batch_size - 1) // batch_size
        
        with state.lock:
            state.running = True
            state.total_chunks = total_chunks
            state.current_chunk = 0
            state.translated_chunks = []
            state.live_feed = ""
            state.start_time = time.time()
            state.current_config = config
        
        # 3. Processing Loop
        for i in range(0, total_chars, batch_size):
            # Check stops
            with state.lock:
                if state.stopped:
                    state.status_message = "Translation stopped."
                    break
                while state.paused:
                    time.sleep(0.5)
            
            chunk = text[i:i+batch_size]
            chunk_num = i // batch_size + 1
            
            with state.lock:
                state.status_message = f"Translating chunk {chunk_num}/{total_chunks}..."
            
            # Translate
            translation = translate_with_ollama(chunk, model=config.get('model'))
            
            with state.lock:
                state.translated_chunks.append(translation)
                state.current_chunk = chunk_num
                state.live_feed += translation + "\n\n"
        
        # 4. Finish
        with state.lock:
            state.running = False
            if not state.stopped:
                state.status_message = "Translation complete."
                # Generate export file immediately
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                # Using English filename for better compatibility, user will see 'translated_...'
                out_filename = f"translated_{ts}.docx"
                out_path = os.path.join(app.config['UPLOAD_FOLDER'], out_filename)
                
                full_text = "\n".join(state.translated_chunks)
                direction = "rtl" if config.get('rtl', True) else "ltr"
                align = "right" if direction == "rtl" else "left"
                
                write_docx(out_path, full_text, direction=direction, alignment=align)
                state.output_file = out_filename
                write_log(f"Saved to {out_path}")

    except Exception as e:
        write_log(f"Worker crashed: {e}")
        with state.lock:
            state.running = False
            state.error = str(e)
            state.status_message = f"Error: {e}"

# ============================================================================
# ROUTES
# ============================================================================

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        return jsonify({'filename': filename})

@app.route('/api/start', methods=['POST'])
def start_translation():
    data = request.json
    filename = data.get('filename')
    if not filename:
        return jsonify({'error': 'Filename required'}), 400
        
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404
    
    state.reset()
    
    thread = threading.Thread(target=worker_thread, args=(filepath, data))
    thread.daemon = True
    thread.start()
    
    return jsonify({'status': 'started'})

@app.route('/api/status', methods=['GET'])
def get_status():
    with state.lock:
        elapsed = 0
        if state.start_time:
            elapsed = int(time.time() - state.start_time)
        
        m, s = divmod(elapsed, 60)
        time_str = f"{m:02d}:{s:02d}"

        return jsonify({
            'running': state.running,
            'paused': state.paused,
            'progress': int((state.current_chunk / state.total_chunks * 100) if state.total_chunks else 0),
            'current_chunk': state.current_chunk,
            'total_chunks': state.total_chunks,
            'elapsed_time': time_str,
            'status_message': state.status_message,
            'live_text': state.live_feed,
            'error': state.error,
            'rtl': state.current_config.get('rtl', True)
        })

@app.route('/api/pause', methods=['POST'])
def pause_translation():
    with state.lock:
        state.paused = not state.paused
        return jsonify({'paused': state.paused})

@app.route('/api/stop', methods=['POST'])
def stop_translation():
    with state.lock:
        state.stopped = True
        state.running = False
    return jsonify({'status': 'stopped'})

@app.route('/api/export', methods=['GET'])
def export_file():
    with state.lock:
        if state.output_file:
            path = os.path.join(app.config['UPLOAD_FOLDER'], state.output_file)
            return send_file(path, as_attachment=True)
        else:
            return "No output file generated yet.", 404


if __name__ == '__main__':
    # Start Flask development server when executed directly
    # Bind to localhost by default for safety; change host if remote access is needed
    app.run(host='127.0.0.1', port=5000, debug=True)
